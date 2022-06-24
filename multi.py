import docx
import os
import re
import pickle
import multiprocessing
from colorama import init
from colorama import Fore
from collections import Counter
from razdel import tokenize
from rnnmorph.predictor import RNNMorphPredictor


def year_record(files, y_start, y_end):
    """Parse year of creation - your files must have it in the last paragraph as mine do"""
    years = {y for y in range(y_start, y_end + 1)}
    results = []
    for file in files:
        doc = docx.Document(file[0])
        lastpara = '\n'.join(p.text for p in list(doc.paragraphs)[-3:])
        for year in years:
            if str(year) in lastpara:
                results.append((*file, year))
                break
    return results


def prons(parse):
    """F*ck RNNMorph who lemmatizes all pronouns to masculine form"""
    if parse.normal_form == "он":
        if "Gender=Fem|Number=Sing" in parse.tag:
            return "она"
        if "Number=Pl" in parse.tag:
            return "они"
        if parse.pos == "DET":
            return parse.word
    return parse.normal_form


def lstsplit(a, n):
    """Split list of paragraphs"""
    k, m = divmod(len(a), n)
    return ['\n'.join(a[i * k + min(i, m):(i + 1) * k + min(i + 1, m)]) for i in range(n)]


def getText(filename):
    """Parse docx file, return string or several strings"""
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    if len(fullText) > 3000:
        return lstsplit(fullText, len(fullText) // 1500)
    return '\n'.join(fullText)


def parsing(lst):
    """Main parsing function. Morphoparsing, returns dict of dicts with frequencies"""
    idee = os.getpid()
    threaddict = dict()
    predictor = RNNMorphPredictor(language="ru")

    for tpl in lst:
        if os.path.getsize(tpl[0]) > 300000:
            print(f'{idee:<6} starts processing {tpl[1]:>35}')
        text = getText(tpl[0])
        if isinstance(text, str):
            text = [text]
        else:
            print(f'{Fore.GREEN}{tpl[1]} has {len(text)} chunks{Fore.RESET}')
        for i in range(len(text)):
            # check for cyrillic words only included
            words = [token.text.lower() for token in tokenize(text[i]) if re.fullmatch(r"(?i)[а-яё.,!?\"«»;:)(—–'’-]+", token.text)]
            pred = predictor.predict(words)
            lemmas = [(prons(word).replace('ё', 'е'), word.pos) for word in pred]
            if tpl[0] in threaddict:
                threaddict[tpl[0]][0]  += Counter(lemmas)
            else:
                threaddict[tpl[0]] = [Counter(lemmas), tpl[2]]
            if i < len(text) - 1:
                print(f'{Fore.GREEN}{idee:<6} parsed chunk {i + 1:>2} of {tpl[1]:>34}{Fore.RESET}')
        print(f'{Fore.CYAN}{idee:<6} parsed {tpl[1]:>46}{Fore.RESET}')
    print(f'{Fore.YELLOW}{idee} finished all parsing, good core{Fore.RESET}')
    return threaddict


if __name__ == '__main__':
    os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3' 
    init()
    filelist = []
    folder = input("Enter folder to parse: ")
    for root, dirs, files in os.walk(folder, topdown=False):
        for name in files:
            if name.endswith('.docx') and '~$' not in name:  # f*ck word's hidden files
                filelist.append((os.path.join(root, name), name[:-5]))
    fileyears = year_record(filelist, 2001, 2022)

    filelists = fileyears[::3], fileyears[1::3], fileyears[2::3]


    pool = multiprocessing.Pool(processes = 3)
    result = pool.map(parsing, filelists)
    bigresult = dict()
    for c in result:
        bigresult.update(c)
  
    pickle.dump(bigresult, open('bigresult.bin', 'wb'))

    totaldict = Counter()
    for key, value in bigresult.items():
        totaldict += value[0]

    with open('totaldict.txt', 'w', encoding='utf8') as textfile:
        longest = max([len(w[0]) for w in totaldict.keys()])
        for key, value in sorted(totaldict.items(), key=lambda x: (-x[1], x[0])):
            if key[1] != 'PUNCT':
                print(f'{key[0]:<{longest}}:{value:>10}', file=textfile)
    
